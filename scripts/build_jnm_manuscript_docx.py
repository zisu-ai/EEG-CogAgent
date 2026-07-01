"""Build the Journal of Neuroscience Methods (JNM) submission DOCX.

This is a thin, JNM-specific adapter over the existing generic manuscript
builder in ``build_manuscript_docx.py``. It reuses every helper from that
module (document styles, title block, reference numbering, table geometry,
markdown-run rendering and figure placement) so the rendered layout stays
identical to the existing narrative-proposal-style manuscript. It only
overrides what differs for the JNM submission:

* source manuscript  -> ``docs/FULL_MANUSCRIPT_JNM.md``
* output              -> ``docs/EEG-CogAgent_JNM_Submission.docx``
* Figure 1 image      -> ``results/ds004504_minimal/figures/figure1_workflow_jnm.png``
* a Supplementary Table S1 (parsed from ``docs/TABLE_JNM_POSITIONING.md``)
  inserted after Table 3 and before the Figure legends. This embedding is the
  default (``include_supplementary_table=True``) so the project master DOCX
  keeps the complete supplementary content; callers that supply the table as a
  separate file pass ``False`` to omit it while retaining the in-text citation.

The generic builder is imported, not altered. Author, affiliation, funding
and repository placeholders are preserved verbatim from the source manuscript
because the user has not supplied those details.
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

# Make the sibling generic builder importable regardless of the CWD.
SCRIPT_DIR = Path(__file__).resolve().parent
if str(SCRIPT_DIR) not in sys.path:
    sys.path.insert(0, str(SCRIPT_DIR))

import build_manuscript_docx as g  # noqa: E402  (reusable, unmodified builder)


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "FULL_MANUSCRIPT_JNM.md"
POSITIONING = ROOT / "docs" / "TABLE_JNM_POSITIONING.md"
OUTPUT = ROOT / "docs" / "EEG-CogAgent_JNM_Submission.docx"

# Figure 1 is the JNM-specific workflow figure; Figures 2-5 match the generic
# builder exactly (same paths).
JNM_FIGURES = {
    1: [ROOT / "results/ds004504_minimal/figures/figure1_workflow_jnm.png"],
    2: [ROOT / "results/ds004504_minimal/figures/figure2_spectral_topomaps.png"],
    3: [ROOT / "results/ds004504_minimal/figures/figure3_connectivity.png"],
    4: [ROOT / "results/ds004504_minimal/figures/figure4_roc.png"],
    5: [ROOT / "results/ds006036_cross_condition/figures/figure5_cross_condition.png"],
}

JNM_FIGURE_ALT = {
    1: (
        "Constrained EEG-CogAgent workflow for Journal of Neuroscience Methods: "
        "BIDS EEG data and a YAML analysis contract drive deterministic MNE-Python "
        "modules for preprocessing, biomarker feature extraction, statistical "
        "testing, machine-learning validation and research outputs. The language-model "
        "layer only plans authorized tools, checks required artifacts and drafts "
        "claim-bounded text; it never diagnoses participants or alters numerical "
        "outputs. The audit contract records participant uniqueness, leakage-safe "
        "predictions, probability bounds, artifact manifests and research-only "
        "claim boundaries."
    ),
    2: "Group-level spectral scalp topographies for delta, theta, alpha, beta and gamma frequency bands.",
    3: "Theta-band weighted phase-lag networks and group comparison of mean network strength.",
    4: "One-versus-rest receiver operating characteristic curves for three nested-cross-validation classifiers.",
    5: "Balanced accuracy and multiclass AUC under participant-disjoint eyes-closed to photomark condition transfer.",
}

# Supplementary Table S1 label (rendered as a Heading 2, consistent with the
# manuscript's own "### Table N." captions) and its fixed DXA column widths.
SUPPLEMENTARY_LABEL = (
    "Supplementary Table S1. Positioning of EEG-CogAgent relative to recent "
    "EEG automation and dementia EEG studies"
)
# 6 columns, fixed DXA geometry, sums to 9360 (= US Letter text width at 1in margins).
S1_WIDTHS = [1400, 1600, 1700, 1450, 1450, 1760]


def prevent_row_split(row) -> None:
    """Keep a table row intact across page boundaries."""
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def parse_md_table(text: str) -> list[list[str]]:
    """Parse a pipe-delimited Markdown table into rows, dropping the separator row."""
    rows: list[list[str]] = []
    for raw in text.splitlines():
        line = raw.strip()
        if not line.startswith("|"):
            continue
        vals = [cell.strip() for cell in line.strip("|").split("|")]
        if all(re.fullmatch(r":?-+:?", cell) for cell in vals):
            continue  # Markdown alignment / separator row
        rows.append(vals)
    return rows


def add_table_with_widths(document, rows: list[list[str]], widths_dxa: list[int]) -> None:
    """Render a Markdown table block with explicit fixed DXA column widths.

    Mirrors the generic builder's ``add_table`` but accepts widths directly so
    the 6-column Supplementary Table S1 can use its own geometry while sharing
    the same shading, repeated header, cell margins and run styling.
    """
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    g.set_table_geometry(table, widths_dxa)
    g.set_repeat_table_header(table.rows[0])
    for r_idx, row in enumerate(rows):
        prevent_row_split(table.rows[r_idx])
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if r_idx == 0:
                g.set_cell_shading(cell, g.LIGHT_FILL)
            paragraph = cell.paragraphs[0]
            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.LEFT if c_idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
            )
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.05
            g.add_md_runs(paragraph, value, base_size=9.0)
            for run in paragraph.runs:
                if r_idx == 0:
                    run.bold = True
    after = document.add_paragraph()
    after.paragraph_format.space_after = Pt(3)


def build(include_supplementary_table: bool = True, out: Path | None = None) -> None:
    """Render the JNM manuscript DOCX.

    ``include_supplementary_table`` controls whether Supplementary Table S1 is
    embedded after Table 3 (default ``True``, matching the project master). When
    ``False`` the in-text citation to Table S1 is retained but the table itself
    is omitted, so a caller that ships the table as a separate editable file
    avoids duplicating it. ``out`` defaults to the project master
    ``docs/EEG-CogAgent_JNM_Submission.docx``; passing a path writes there
    instead, leaving the master untouched.
    """
    text = SOURCE.read_text(encoding="utf-8")
    lines = text.splitlines()
    positioning_rows = (
        parse_md_table(POSITIONING.read_text(encoding="utf-8"))
        if include_supplementary_table
        else []
    )

    # Adapt the shared builder's figure registry in place (it is read by
    # ``g.add_figure``). Figure 1 path and alt text are JNM-specific; the rest
    # is identical to the generic builder.
    g.FIGURES = JNM_FIGURES
    g.FIGURE_ALT = JNM_FIGURE_ALT

    doc = Document()
    g.configure_styles(doc)
    doc.settings.odd_and_even_pages_header_footer = False
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    # Academic-manuscript override: no running header; retain page numbers only.
    section.header.paragraphs[0].text = ""
    g.add_page_number(section.footer.paragraphs[0])

    # Title block: title, author/affiliation/correspondence placeholders.
    title_end = next(i for i, line in enumerate(lines) if line == "## Abstract")
    g.add_title_block(doc, lines[:title_end])

    ref_num_id = g.add_reference_numbering(doc)
    in_references = False
    in_figures = False
    table_idx = 0
    i = title_end

    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue

        if line.startswith("## "):
            heading = line[3:].strip()
            # Insert Supplementary Table S1 immediately before the Figure
            # legends, i.e. after Table 3 and before the figures. It is
            # rendered as a labeled table, not as prose. Omitted when the
            # caller supplies it as a separate file; the in-text citation
            # in the prose is retained either way.
            if heading == "Figure legends" and include_supplementary_table:
                sup = doc.add_paragraph(style="Heading 2")
                sup.alignment = WD_ALIGN_PARAGRAPH.LEFT
                g.add_md_runs(sup, SUPPLEMENTARY_LABEL, base_size=13, base_color=g.BLUE)
                add_table_with_widths(doc, positioning_rows, S1_WIDTHS)
            in_references = heading == "References"
            in_figures = heading == "Figure legends"
            p = doc.add_paragraph(style="Heading 1")
            if heading == "Figure legends":
                p.paragraph_format.page_break_before = True
            g.add_md_runs(p, heading, base_size=16, base_color=g.BLUE)
            i += 1
            continue

        if line.startswith("### "):
            heading = line[4:].strip()
            p = doc.add_paragraph(style="Heading 2")
            if heading.startswith("Table 3."):
                p.paragraph_format.page_break_before = True
            g.add_md_runs(p, heading, base_size=13, base_color=g.BLUE)
            if heading.startswith("Table "):
                table_idx += 1
            i += 1
            continue

        if line.startswith("# "):
            i += 1
            continue

        if line.startswith("|"):
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                vals = [x.strip() for x in lines[i].strip().strip("|").split("|")]
                if not all(re.fullmatch(r":?-+:?", x) for x in vals):
                    rows.append(vals)
                i += 1
            g.add_table(doc, rows, table_idx)
            continue

        if in_references and re.match(r"^\d+\.\s", line):
            content = re.sub(r"^\d+\.\s*", "", line)
            p = doc.add_paragraph()
            g.apply_num(p, ref_num_id)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(5)
            p.paragraph_format.line_spacing = 1.15
            g.add_md_runs(p, content, base_size=9.5)
            i += 1
            continue

        fig_match = re.match(r"^\*\*Figure (\d+)\.", line)
        if in_figures and fig_match:
            number = int(fig_match.group(1))
            g.add_figure(doc, number)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(10)
            p.paragraph_format.keep_with_next = False
            g.add_md_runs(p, line, base_size=9.3)
            i += 1
            continue

        # Join consecutive prose lines, although the Markdown source uses one line per paragraph.
        prose = [line]
        i += 1
        while i < len(lines):
            nxt = lines[i].strip()
            if not nxt or nxt.startswith(("#", "|")) or (in_references and re.match(r"^\d+\.\s", nxt)):
                break
            prose.append(nxt)
            i += 1
        paragraph_text = " ".join(prose)
        p = doc.add_paragraph()
        p.paragraph_format.keep_together = False
        p.paragraph_format.widow_control = True
        if paragraph_text.startswith("**Keywords:"):
            p.paragraph_format.space_after = Pt(12)
        g.add_md_runs(p, paragraph_text)

    # The shared table builder permits row splitting. JNM tables contain
    # multi-line evidence statements, so keep every row intact for clean page
    # transitions while retaining automatic row height.
    for table in doc.tables:
        for row in table.rows:
            prevent_row_split(row)

    # Document metadata. Placeholders are retained and nothing is fabricated.
    core = doc.core_properties
    core.title = (
        "EEG-CogAgent: An Auditable Language-Model-Assisted Workflow for "
        "Reproducible Dementia EEG Biomarker Discovery"
    )
    core.subject = "Journal of Neuroscience Methods submission manuscript"
    core.keywords = (
        "EEG; dementia; language model; agent; Alzheimer's disease; "
        "frontotemporal dementia; biomarkers; reproducibility; BIDS"
    )
    core.author = "[Author names]"
    core.comments = (
        "JNM submission built from FULL_MANUSCRIPT_JNM.md. Author, affiliation, "
        "funding and repository placeholders are retained pending user input."
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    output_path = out if out is not None else OUTPUT
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(output_path)


if __name__ == "__main__":
    build()
