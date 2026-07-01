from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "FULL_MANUSCRIPT.md"
OUTPUT = ROOT / "docs" / "EEG-CogAgent_Complete_Manuscript.docx"

FIGURES = {
    1: [ROOT / "results/ds004504_minimal/figures/figure1_workflow.png"],
    2: [ROOT / "results/ds004504_minimal/figures/figure2_spectral_topomaps.png"],
    3: [ROOT / "results/ds004504_minimal/figures/figure3_connectivity.png"],
    4: [ROOT / "results/ds004504_minimal/figures/figure4_roc.png"],
    5: [ROOT / "results/ds006036_cross_condition/figures/figure5_cross_condition.png"],
}

FIGURE_ALT = {
    1: "EEG-CogAgent workflow from BIDS input through deterministic analysis, validation, audit and report generation.",
    2: "Group-level spectral scalp topographies for delta, theta, alpha, beta and gamma frequency bands.",
    3: "Theta-band weighted phase-lag networks and group comparison of mean network strength.",
    4: "One-versus-rest receiver operating characteristic curves for three nested-cross-validation classifiers.",
    5: "Balanced accuracy and multiclass AUC under participant-disjoint eyes-closed to photomark condition transfer.",
}

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
GRAY = "666666"
LIGHT_FILL = "F4F6F9"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = 120) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[idx]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            set_cell_margins(cell)


def set_font(run, name="Calibri", size=11, bold=None, italic=None, color=None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_plain_runs(paragraph, text: str, base_size=11, base_color=None, italic_all=False) -> None:
    exponent_re = re.compile(r"(10\^(-?\d+))")
    pos = 0
    for match in exponent_re.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos:match.start()])
            set_font(run, size=base_size, italic=italic_all, color=base_color)
        base = paragraph.add_run("10")
        set_font(base, size=base_size, italic=italic_all, color=base_color)
        exponent = paragraph.add_run(match.group(2))
        set_font(exponent, size=max(base_size - 1.5, 7), italic=italic_all, color=base_color)
        exponent.font.superscript = True
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_font(run, size=base_size, italic=italic_all, color=base_color)


def add_md_runs(paragraph, text: str, base_size=11, base_color=None, italic_all=False) -> None:
    token_re = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")
    pos = 0
    for match in token_re.finditer(text):
        if match.start() > pos:
            add_plain_runs(paragraph, text[pos:match.start()], base_size, base_color, italic_all)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_font(run, size=base_size, bold=True, italic=italic_all, color=base_color)
        elif token.startswith("*"):
            run = paragraph.add_run(token[1:-1])
            set_font(run, size=base_size, italic=True, color=base_color)
        else:
            run = paragraph.add_run(token[1:-1])
            set_font(run, name="Consolas", size=max(base_size - 1, 8), color=base_color)
        pos = match.end()
    if pos < len(text):
        add_plain_runs(paragraph, text[pos:], base_size, base_color, italic_all)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2])
    set_font(run, size=9, color=GRAY)


def add_reference_numbering(document: Document) -> int:
    numbering = document.part.numbering_part.element
    abstract_num = OxmlElement("w:abstractNum")
    abstract_num.set(qn("w:abstractNumId"), "42")
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1.")
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "space")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "360")
    p_pr.extend([tabs, ind])
    lvl.extend([start, num_fmt, lvl_text, suff, p_pr])
    abstract_num.append(lvl)
    numbering.append(abstract_num)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), "42")
    abs_id = OxmlElement("w:abstractNumId")
    abs_id.set(qn("w:val"), "42")
    num.append(abs_id)
    numbering.append(num)
    return 42


def apply_num(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_el])
    p_pr.append(num_pr)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333
    normal.paragraph_format.widow_control = True

    style_specs = {
        "Title": (20, "1F1F1F", 0, 10),
        "Subtitle": (11, GRAY, 0, 4),
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in style_specs.items():
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_title_block(doc: Document, lines: list[str]) -> None:
    title = lines[0][2:].strip()
    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_md_runs(p, title, base_size=20)

    for raw in lines[1:]:
        if not raw.strip():
            continue
        p = doc.add_paragraph(style="Subtitle")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_md_runs(p, raw, base_size=10.5, base_color=GRAY)
    rule = doc.add_paragraph()
    rule.paragraph_format.space_before = Pt(4)
    rule.paragraph_format.space_after = Pt(12)
    p_pr = rule._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), BLUE)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_table(doc: Document, rows: list[list[str]], table_index: int) -> None:
    widths = {
        1: [1730, 850, 2450, 2200, 2130],
        2: [3160, 1730, 2300, 2170],
        3: [2020, 3670, 3670],
    }[table_index]
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    set_repeat_table_header(table.rows[0])
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if r_idx == 0:
                set_cell_shading(cell, LIGHT_FILL)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if c_idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            add_md_runs(p, value, base_size=9.2)
            for run in p.runs:
                if r_idx == 0:
                    run.bold = True
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(3)


def add_figure(doc: Document, number: int) -> None:
    paths = FIGURES[number]
    if number == 2 and len(paths) == 5:
        labels = ["A", "B", "C", "D", "E"]
        for row_indices in ((0, 1, 2), (3, 4)):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.keep_with_next = True
            for idx in row_indices:
                label_run = p.add_run(labels[idx] + " ")
                set_font(label_run, size=8, bold=True)
                shape = p.add_run().add_picture(str(paths[idx]), width=Inches(1.82))
                band = paths[idx].stem.replace("topomap_", "")
                shape._inline.docPr.set("descr", f"{band.capitalize()}-band scalp topographies for AD, FTD and healthy controls.")
                shape._inline.docPr.set("title", f"Figure 2{labels[idx]} {band} topographies")
                if idx != row_indices[-1]:
                    p.add_run("  ")
    else:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.keep_with_next = True
        shape = p.add_run().add_picture(str(paths[0]), width=Inches(6.15))
        shape._inline.docPr.set("descr", FIGURE_ALT[number])
        shape._inline.docPr.set("title", f"Figure {number}")


def build() -> None:
    text = SOURCE.read_text(encoding="utf-8")
    lines = text.splitlines()
    doc = Document()
    configure_styles(doc)
    doc.settings.odd_and_even_pages_header_footer = False
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    # Academic-manuscript override: no running header; retain page numbers only.
    header = section.header.paragraphs[0]
    header.text = ""
    add_page_number(section.footer.paragraphs[0])

    # First block: title, author placeholders, affiliations, correspondence.
    title_end = next(i for i, line in enumerate(lines) if line == "## Abstract")
    add_title_block(doc, lines[:title_end])

    ref_num_id = add_reference_numbering(doc)
    in_references = False
    in_tables = False
    in_figures = False
    table_idx = 0
    i = title_end

    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue

        if line.startswith("## "):
            heading = line[3:].strip()
            in_references = heading == "References"
            in_tables = heading == "Tables"
            in_figures = heading == "Figure legends"
            p = doc.add_paragraph(style="Heading 1")
            add_md_runs(p, heading, base_size=16, base_color=BLUE)
            i += 1
            continue

        if line.startswith("### "):
            heading = line[4:].strip()
            p = doc.add_paragraph(style="Heading 2")
            add_md_runs(p, heading, base_size=13, base_color=BLUE)
            if heading.startswith("Table "):
                table_idx += 1
            i += 1
            continue

        if line.startswith("# "):
            i += 1
            continue

        if line.startswith("|"):
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                vals = [x.strip() for x in lines[i].strip().strip("|").split("|")]
                if not all(re.fullmatch(r":?-+:?", x) for x in vals):
                    rows.append(vals)
                i += 1
            add_table(doc, rows, table_idx)
            continue

        if in_references and re.match(r"^\d+\.\s", line):
            content = re.sub(r"^\d+\.\s*", "", line)
            p = doc.add_paragraph()
            apply_num(p, ref_num_id)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(5)
            p.paragraph_format.line_spacing = 1.15
            add_md_runs(p, content, base_size=9.5)
            i += 1
            continue

        fig_match = re.match(r"^\*\*Figure (\d+)\.", line)
        if in_figures and fig_match:
            number = int(fig_match.group(1))
            add_figure(doc, number)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(10)
            p.paragraph_format.keep_with_next = False
            add_md_runs(p, line, base_size=9.3)
            i += 1
            continue

        # Join consecutive prose lines, although the Markdown source uses one line per paragraph.
        prose = [line]
        i += 1
        while i < len(lines):
            nxt = lines[i].strip()
            if not nxt or nxt.startswith(("#", "|")) or (in_references and re.match(r"^\d+\.\s", nxt)):
                break
            prose.append(nxt)
            i += 1
        paragraph_text = " ".join(prose)
        p = doc.add_paragraph()
        p.paragraph_format.keep_together = False
        p.paragraph_format.widow_control = True
        if paragraph_text.startswith("**Keywords:"):
            p.paragraph_format.space_after = Pt(12)
        add_md_runs(p, paragraph_text)

    core = doc.core_properties
    core.title = "EEG-CogAgent: An Auditable Language-Model Agent for Reproducible Dementia EEG Biomarker Analysis"
    core.subject = "Complete scientific manuscript"
    core.keywords = "EEG, dementia, language model, reproducibility, BIDS"
    core.author = "[Author names]"
    core.comments = "Generated from the verified project manuscript."

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
