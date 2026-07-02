"""Build the v3.1 JNM manuscript DOCX from docs/FULL_MANUSCRIPT_JNM_V3_1.md.

Differences from the v3 JNM builder (build_jnm_manuscript_docx.py):
* source            -> docs/FULL_MANUSCRIPT_JNM_V3_1.md
* references        -> author-year, alphabetical (rendered as hanging-indent
  paragraphs, NOT numbered); plus [dataset] entries
* figures           -> Figures 1-6 (Figure 6 is the new external-integrity composite)
* tables            -> Tables 1-4 (Table 4 is the new external-evaluation table)

Reuses every reusable helper from build_manuscript_docx.py (g). No scientific
result is altered; placeholders for author/affiliation/funding/repository are
preserved verbatim.
"""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

SCRIPT_DIR = Path(__file__).resolve().parent
if str(SCRIPT_DIR) not in sys.path:
    sys.path.insert(0, str(SCRIPT_DIR))

import build_manuscript_docx as g  # noqa: E402

# Fixed DXA column geometry for tables 1-3 (matches the generic builder). Table 4
# (new external-evaluation table) and any future table fall back to even widths.
_KNOWN_TABLE_WIDTHS = {
    1: [1730, 850, 2450, 2200, 2130],
    2: [3160, 1730, 2300, 2170],
    3: [2020, 3670, 3670],
}


def _table_widths(rows: list[list[str]], table_idx: int) -> list[int]:
    ncol = len(rows[0])
    known = _KNOWN_TABLE_WIDTHS.get(table_idx)
    if known and len(known) == ncol:
        return known
    per = 9360 // ncol
    widths = [per] * ncol
    widths[-1] = 9360 - per * (ncol - 1)
    return widths


def _add_table(doc: Document, rows: list[list[str]], table_idx: int) -> None:
    widths = _table_widths(rows, table_idx)
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    g.set_table_geometry(table, widths)
    g.set_repeat_table_header(table.rows[0])
    for r_idx, row in enumerate(rows):
        _prevent_row_split(table.rows[r_idx])
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if r_idx == 0:
                g.set_cell_shading(cell, g.LIGHT_FILL)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if c_idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            g.add_md_runs(p, value, base_size=9.2)
            for run in p.runs:
                if r_idx == 0:
                    run.bold = True
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(3)

ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "FULL_MANUSCRIPT_JNM_V3_1.md"
OUTPUT_DEFAULT = ROOT / "docs" / "EEG-CogAgent_JNM_Submission_V3_1.docx"

DISC = ROOT / "results" / "ds004504_minimal" / "figures"
XCOND = ROOT / "results" / "ds006036_cross_condition" / "figures"
EXT = ROOT / "results" / "external_validation_osf_v3_1" / "figures"

V31_FIGURES = {
    1: [DISC / "figure1_workflow_jnm.png"],
    2: [DISC / "figure2_spectral_topomaps.png"],
    3: [DISC / "figure3_connectivity.png"],
    4: [DISC / "figure4_roc.png"],
    5: [XCOND / "figure5_cross_condition.png"],
    6: [EXT / "figure6_external_integrity.png"],
}
V31_FIGURE_ALT = {
    1: "Constrained EEG-CogAgent workflow: BIDS EEG and a YAML contract drive deterministic MNE-Python modules; the language-model layer plans, checks artifacts, and drafts claim-bounded text but never diagnoses or alters numbers; audit gates include content-level fingerprints, leakage-safe predictions, self-verified manifests, and an independent external evaluation branch.",
    2: "Group-level spectral scalp topographies for delta, theta, alpha, beta and gamma bands.",
    3: "Functional connectivity and network findings: lower alpha coherence and higher theta weighted-phase-lag-index coupling in dementia versus controls.",
    4: "One-versus-rest receiver operating characteristic curves for the three discovery nested-cross-validation classifiers.",
    5: "Participant-disjoint cross-condition transfer from eyes-closed to photomark recordings within the same cohort (not external validation).",
    6: "Independent external archive (OSF node 2v5md): integrity flow from 92 nominal records to 88 unique recordings, primary unique-record ROC, confusion matrix, and top domain-shift Cohen's d.",
}


def _prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def _render_reference(paragraph, text: str) -> None:
    """Author-year reference as a hanging-indent paragraph."""
    paragraph.paragraph_format.left_indent = Pt(18)
    paragraph.paragraph_format.first_line_indent = Pt(-18)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.12
    g.add_md_runs(paragraph, text, base_size=9.5)


def build(out: Path | None = None) -> Path:
    text = SOURCE.read_text(encoding="utf-8")
    lines = text.splitlines()

    g.FIGURES = V31_FIGURES
    g.FIGURE_ALT = V31_FIGURE_ALT

    doc = Document()
    g.configure_styles(doc)
    doc.settings.odd_and_even_pages_header_footer = False
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header.paragraphs[0].text = ""
    g.add_page_number(section.footer.paragraphs[0])

    title_end = next(i for i, line in enumerate(lines) if line == "## Abstract")
    g.add_title_block(doc, lines[:title_end])

    in_references = False
    in_figures = False
    table_idx = 0
    i = title_end

    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue

        if line.startswith("## "):
            heading = line[3:].strip()
            in_references = heading == "References"
            in_figures = heading == "Figure legends"
            p = doc.add_paragraph(style="Heading 1")
            if heading == "Figure legends":
                p.paragraph_format.page_break_before = True
            g.add_md_runs(p, heading, base_size=16, base_color=g.BLUE)
            i += 1
            continue

        if line.startswith("### "):
            heading = line[4:].strip()
            p = doc.add_paragraph(style="Heading 2")
            if heading.startswith("Table 3.") or heading.startswith("Table 4."):
                p.paragraph_format.page_break_before = True
            g.add_md_runs(p, heading, base_size=13, base_color=g.BLUE)
            if heading.startswith("Table "):
                table_idx += 1
            i += 1
            continue

        if line.startswith("# "):
            i += 1
            continue

        if line.startswith("|"):
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                vals = [x.strip() for x in lines[i].strip().strip("|").split("|")]
                if not all(re.fullmatch(r":?-+:?", x) for x in vals):
                    rows.append(vals)
                i += 1
            _add_table(doc, rows, table_idx)
            continue

        if in_references:
            _render_reference(doc.add_paragraph(), line)
            i += 1
            continue

        fig_match = re.match(r"^\*\*Figure (\d+)\.", line)
        if in_figures and fig_match:
            number = int(fig_match.group(1))
            g.add_figure(doc, number)
            # Force each figure (except Figure 1, which follows the page-broken
            # "Figure legends" heading) onto a fresh page so the inline image and
            # its caption stay together; otherwise large composites clip across
            # the page boundary and the caption separates from the figure.
            if number >= 2:
                target = doc.paragraphs[-2] if number == 2 else doc.paragraphs[-1]
                target.paragraph_format.page_break_before = True
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(10)
            g.add_md_runs(p, line, base_size=9.3)
            i += 1
            continue

        prose = [line]
        i += 1
        while i < len(lines):
            nxt = lines[i].strip()
            if not nxt or nxt.startswith(("#", "|")):
                break
            prose.append(nxt)
            i += 1
        paragraph_text = " ".join(prose)
        p = doc.add_paragraph()
        p.paragraph_format.widow_control = True
        if paragraph_text.startswith("**Keywords:"):
            p.paragraph_format.space_after = Pt(12)
        g.add_md_runs(p, paragraph_text)

    for table in doc.tables:
        for row in table.rows:
            _prevent_row_split(row)

    core = doc.core_properties
    core.title = (
        "EEG-CogAgent: An Auditable Language-Model-Assisted Workflow with "
        "Content-Level Integrity Audit and Independent External Evaluation of "
        "Reproducible Dementia EEG Biomarkers"
    )
    core.subject = "Journal of Neuroscience Methods submission manuscript (v3.1)"
    core.keywords = ("EEG; dementia; language model; agent; Alzheimer's disease; "
                     "frontotemporal dementia; biomarkers; reproducibility; audit; BIDS")
    core.author = "[Author names]"
    core.comments = ("JNM v3.1 submission built from FULL_MANUSCRIPT_JNM_V3_1.md. "
                     "Author/affiliation/funding/repository placeholders retained.")

    output_path = out if out is not None else OUTPUT_DEFAULT
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(output_path)
    return output_path


if __name__ == "__main__":
    build()
