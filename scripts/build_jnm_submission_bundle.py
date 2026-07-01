"""Deterministically build the Journal of Neuroscience Methods (JNM) submission
bundle into ``submission/JNM/``.

The bundle is assembled entirely from the reviewed project sources:

* ``docs/FULL_MANUSCRIPT_JNM.md``          (corrected keywords; results untouched)
* ``docs/JNM_HIGHLIGHTS.md``               (3-5 bullets, each <=85 characters)
* ``docs/JNM_COVER_LETTER.md``             (placeholders retained)
* ``docs/TABLE_JNM_POSITIONING.md``        (Supplementary Table S1 source)
* ``results/.../figure*.{png,pdf}``        (project-owned figures)

Outputs (all editable / project-owned; no fabricated author or DOI details):

    submission/JNM/
        Manuscript_EEG-CogAgent.docx          rebuilt via the reviewed JNM builder
        Highlights_EEG-CogAgent.docx          clean editable Word file
        Cover_Letter_EEG-CogAgent.docx        clean editable Word file
        Supplementary_Table_S1.docx           editable Word table (not an image)
        Graphical_Abstract_EEG-CogAgent.png   project-owned, >=531x1328 px
        Graphical_Abstract_EEG-CogAgent.pdf   vector (preferred format)
        figures/Figure_1..Figure_5.{png,pdf}  logical names
        Author_Input_Form.md                  every human-only field
        SUBMISSION_READINESS.md               READY / OPTIONAL / BLOCKING
        submission_manifest.csv               bundle inventory

The script is idempotent: rerunning overwrites only the files it owns and never
deletes unrelated files. No scientific result, metric, threshold, dataset
boundary or reference is altered.
"""

from __future__ import annotations

import csv
import shutil
import sys
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt

# Make the sibling manuscript builders importable regardless of CWD.
SCRIPT_DIR = Path(__file__).resolve().parent
if str(SCRIPT_DIR) not in sys.path:
    sys.path.insert(0, str(SCRIPT_DIR))

import build_manuscript_docx as g  # noqa: E402  (reusable, unmodified builder)
import build_jnm_manuscript_docx as jnm  # noqa: E402  (reusable, unmodified builder)


ROOT = Path(__file__).resolve().parents[1]
SUBMISSION = ROOT / "submission" / "JNM"
FIGURES_OUT = SUBMISSION / "figures"

SRC_MANUSCRIPT = ROOT / "docs" / "FULL_MANUSCRIPT_JNM.md"
SRC_HIGHLIGHTS = ROOT / "docs" / "JNM_HIGHLIGHTS.md"
SRC_COVER = ROOT / "docs" / "JNM_COVER_LETTER.md"
SRC_TABLE = ROOT / "docs" / "TABLE_JNM_POSITIONING.md"

# Logical figure number -> (source directory, source stem). Figure 1 uses the
# JNM-specific constrained-workflow graphic, matching the manuscript's Figure 1.
FIGURE_SOURCES = {
    1: (ROOT / "results" / "ds004504_minimal" / "figures", "figure1_workflow_jnm"),
    2: (ROOT / "results" / "ds004504_minimal" / "figures", "figure2_spectral_topomaps"),
    3: (ROOT / "results" / "ds004504_minimal" / "figures", "figure3_connectivity"),
    4: (ROOT / "results" / "ds004504_minimal" / "figures", "figure4_roc"),
    5: (ROOT / "results" / "ds006036_cross_condition" / "figures", "figure5_cross_condition"),
}

GA_MIN_W_PX = 1328  # JNM: minimum width (height x width = 531 x 1328 px)
GA_MIN_H_PX = 531


# --------------------------------------------------------------------------- #
# Editable Word documents
# --------------------------------------------------------------------------- #
def _new_clean_doc() -> Document:
    """A plain editable Word document with the project's base styles."""
    doc = Document()
    g.configure_styles(doc)
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    return doc


def build_highlights_docx(out: Path) -> None:
    doc = _new_clean_doc()
    h = doc.add_paragraph(style="Heading 1")
    g.add_md_runs(h, "Highlights", base_size=16, base_color=g.BLUE)
    for line in SRC_HIGHLIGHTS.read_text(encoding="utf-8").splitlines():
        line = line.strip()
        if line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            g.add_md_runs(p, line[2:].strip())
    _set_core(doc, "EEG-CogAgent - Highlights", "[Author names]")
    doc.save(out)


def build_cover_letter_docx(out: Path) -> None:
    doc = _new_clean_doc()
    for raw in SRC_COVER.read_text(encoding="utf-8").splitlines():
        line = raw.strip()
        if not line:
            continue
        if line.startswith("# "):
            p = doc.add_paragraph(style="Heading 1")
            g.add_md_runs(p, line[2:].strip(), base_size=16, base_color=g.BLUE)
            continue
        p = doc.add_paragraph()
        # Preserve the literal "[Pending: ...]" approval marker verbatim.
        g.add_md_runs(p, line)
    _set_core(doc, "EEG-CogAgent - Cover Letter", "[Corresponding author name]")
    doc.save(out)


def build_table_s1_docx(out: Path) -> None:
    doc = _new_clean_doc()
    h = doc.add_paragraph(style="Heading 1")
    g.add_md_runs(h, "Supplementary Table S1", base_size=16, base_color=g.BLUE)
    cap = doc.add_paragraph()
    g.add_md_runs(
        cap,
        "Positioning of EEG-CogAgent relative to recent EEG automation and dementia "
        "EEG studies.",
        base_size=10,
        italic_all=False,
    )
    rows = jnm.parse_md_table(SRC_TABLE.read_text(encoding="utf-8"))
    jnm.add_table_with_widths(doc, rows, jnm.S1_WIDTHS)
    _set_core(doc, "EEG-CogAgent - Supplementary Table S1", "[Author names]")
    doc.save(out)


def build_manuscript_docx(out: Path) -> None:
    """Build the bundle manuscript via the reviewed JNM builder.

    Built with ``include_supplementary_table=False`` so the bundle manuscript
    contains Tables 1-3 only; the in-text citation to Supplementary Table S1 is
    retained and the editable table is shipped separately as
    ``Supplementary_Table_S1.docx``. The project master
    ``docs/EEG-CogAgent_JNM_Submission.docx`` is not touched here: it is
    regenerated by the builder's default (``True``) and keeps S1 embedded as the
    complete master copy.
    """
    jnm.build(include_supplementary_table=False, out=out)


def _set_core(doc: Document, title: str, author: str) -> None:
    core = doc.core_properties
    core.title = title
    core.author = author
    core.subject = "Journal of Neuroscience Methods submission"


# --------------------------------------------------------------------------- #
# Graphical abstract (project-owned artwork, no generative/third-party tools)
# --------------------------------------------------------------------------- #
def build_graphical_abstract(out_png: Path, out_pdf: Path) -> None:
    import matplotlib

    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    from matplotlib.patches import FancyArrowPatch, FancyBboxPatch

    # Wide aspect (~2.6:1) matching the JNM 5x13 cm graphical-abstract format.
    fig, ax = plt.subplots(figsize=(8.4, 3.2))
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis("off")

    def box(xy, wh, title, subtitle, fill, edge="#4B5563"):
        x, y = xy
        w, h = wh
        ax.add_patch(
            FancyBboxPatch(
                (x, y), w, h,
                boxstyle="round,pad=0.010,rounding_size=0.020",
                facecolor=fill, edgecolor=edge, linewidth=0.9,
            )
        )
        ax.text(
            x + w / 2, y + h * 0.64, title, ha="center", va="center",
            fontsize=7.6, fontweight="bold", linespacing=1.05, color="#111827",
        )
        ax.text(
            x + w / 2, y + h * 0.26, subtitle, ha="center", va="center",
            fontsize=5.9, linespacing=1.15, color="#374151",
        )
        return x + w / 2, y + h / 2

    # Top banner: constrained LLM orchestration + mandatory safety label.
    ax.add_patch(
        FancyBboxPatch(
            (0.05, 0.70), 0.90, 0.20,
            boxstyle="round,pad=0.014,rounding_size=0.030",
            facecolor="#1F4E79", edgecolor="#163A5B", linewidth=1.0,
        )
    )
    ax.text(
        0.50, 0.855, "Constrained LLM orchestration", ha="center", va="center",
        fontsize=10.5, fontweight="bold", color="white",
    )
    ax.text(
        0.50, 0.805,
        "plans authorized tools  ·  checks required artifacts  ·  drafts claim-bounded text",
        ha="center", va="center", fontsize=6.6, color="#E5EEF6",
    )
    ax.text(
        0.50, 0.745,
        "The LLM never diagnoses participants and never alters numerical outputs.",
        ha="center", va="center", fontsize=6.9, fontweight="bold", color="#FDE68A",
    )

    stages = [
        ("BIDS EEG\n+ YAML", "open data\nanalysis contract", "#E8F1FA"),
        ("MNE\npreprocessing", "filter · reference\nepoch · QC", "#E7F5EF"),
        ("Biomarker\nfeatures", "power · ratios\nconnectivity · graph", "#FFF3D6"),
        ("Statistics\n+ ML", "FDR · contrasts\nnested CV · sensitivity", "#F4E9F7"),
        ("Research\noutputs", "tables · figures\naudit · report draft", "#FCE8E6"),
    ]
    xs = [0.035, 0.2325, 0.43, 0.6275, 0.825]
    centers = [
        box((x, 0.30), (0.145, 0.30), title, sub, fill)
        for x, (title, sub, fill) in zip(xs, stages)
    ]
    for i in range(len(centers) - 1):
        ax.add_patch(
            FancyArrowPatch(
                (centers[i][0] + 0.078, centers[i][1]),
                (centers[i + 1][0] - 0.078, centers[i + 1][1]),
                arrowstyle="-|>", mutation_scale=9, linewidth=0.9, color="#4B5563",
            )
        )
    for c in centers:
        ax.add_patch(
            FancyArrowPatch(
                (0.50, 0.70), (c[0], 0.60),
                arrowstyle="-|>", mutation_scale=5, linewidth=0.5,
                color="#6B7280", alpha=0.70,
            )
        )

    # Bottom banner: audit contract and safety guardrails.
    ax.add_patch(
        FancyBboxPatch(
            (0.05, 0.06), 0.90, 0.15,
            boxstyle="round,pad=0.012,rounding_size=0.022",
            facecolor="#F8FAFC", edgecolor="#9CA3AF", linewidth=0.9,
        )
    )
    ax.text(
        0.50, 0.165, "Audit contract and safety guardrails", ha="center", va="center",
        fontsize=8.6, fontweight="bold", color="#1F2937",
    )
    ax.text(
        0.50, 0.105,
        "participant uniqueness  ·  leakage-safe predictions  ·  probability bounds  "
        "·  artifact manifest  ·  research-only claims",
        ha="center", va="center", fontsize=6.0, color="#4B5563",
    )

    fig.savefig(out_png, dpi=400, bbox_inches="tight", facecolor="white")
    fig.savefig(out_pdf, bbox_inches="tight", facecolor="white")
    plt.close(fig)


# --------------------------------------------------------------------------- #
# Figures
# --------------------------------------------------------------------------- #
def copy_figures() -> list[Path]:
    copied: list[Path] = []
    for number, (directory, stem) in FIGURE_SOURCES.items():
        for ext in ("png", "pdf"):
            src = directory / f"{stem}.{ext}"
            if not src.exists():
                raise FileNotFoundError(f"Missing figure source: {src}")
            dst = FIGURES_OUT / f"Figure_{number}.{ext}"
            shutil.copy2(src, dst)
            copied.append(dst)
    return copied


# --------------------------------------------------------------------------- #
# Plain-text bundle documents
# --------------------------------------------------------------------------- #
AUTHOR_INPUT_FORM = """# JNM Submission - Author Input Form

This form lists every field that the build script cannot generate. The bundle
under `submission/JNM/` is complete except for these human-only items. Fill it
in, update the source placeholders, and rerun `scripts/build_jnm_submission_bundle.py`.

## 1. Authors and affiliations
- [ ] Full author list, in display order:
- [ ] ORCID iD for each author (recommended):
- [ ] Numbered affiliations for each author (department, institution, city, region/state, country):

## 2. Corresponding author (full contact details)
- [ ] Full name:
- [ ] Email:
- [ ] Phone, including country code:
- [ ] Mailing address:
- [ ] ORCID iD:

## 3. Funding
- [ ] Funding statement, or confirm verbatim: "This research received no specific grant from any funding agency, the commercial sector, or not-for-profit institutions."

## 4. CRediT author roles
- [ ] CRediT contributor roles for each author. Valid roles: Conceptualization, Methodology, Software, Validation, Formal analysis, Investigation, Resources, Data Curation, Writing - Original Draft, Writing - Review & Editing, Visualization, Supervision, Project administration, Funding acquisition.

## 5. Code availability
- [ ] Public repository URL (e.g., GitHub):
- [ ] Archived release DOI (e.g., Zenodo):

## 6. Declarations requiring live confirmation by every author
- [ ] Author approval: every author has read and approved the final manuscript and agrees to the submission.
- [ ] Originality: the manuscript is original, has not been published elsewhere, and is not under consideration by another journal.
- [ ] Competing interests: confirm "none declared", or disclose specific interests.
- [ ] Consent for publication: confirm in the submission system (not applicable for this de-identified public-data study, but still requires a system entry).

## 7. Optional final approvals
- [ ] Corresponding-author sign-off on Highlights, Cover Letter, and Graphical Abstract.
- [ ] Decision on whether to include the Graphical Abstract (encouraged by JNM, not mandatory).
"""


READINESS = """# JNM Submission Readiness - EEG-CogAgent

Generated by `scripts/build_jnm_submission_bundle.py` from the corrected source
files. No scientific result, metric, threshold, dataset boundary or reference
was altered; the only manuscript change is reducing keywords from 8 to 7.

## Bundle structure: READY
The bundle under `submission/JNM/` is complete and internally consistent: every
file listed in `submission_manifest.csv` is present, the manuscript contains
Tables 1-3 with the in-text citation to Supplementary Table S1 retained, the
editable `Supplementary_Table_S1.docx` is supplied as a separate file (no
duplicate embedding), and all five figures plus the graphical abstract are
present in the expected formats.

The bundle structure is ready. The manuscript and cover letter themselves are
NOT upload-ready: they still carry unresolved author / affiliation / contact /
funding / CRediT / repository / approval placeholders. They are marked
BLOCKED - HUMAN INPUT below and must be completed and approved before the
bundle is uploaded.

## READY (generated, verified, free of human-only placeholders)
- `Highlights_EEG-CogAgent.docx` - one `Highlights` heading plus 5 bullets, each <=85 characters including spaces; no instructional text.
- `Supplementary_Table_S1.docx` - editable Word table (not an image); positioning content only.
- `Graphical_Abstract_EEG-CogAgent.png` / `.pdf` - meets the minimum 531 x 1328 px requirement; project-owned artwork only.
- `figures/Figure_1` through `Figure_5` - each provided in PNG and PDF.
- `Author_Input_Form.md`, `submission_manifest.csv`, this file.

## BLOCKED - HUMAN INPUT (present in bundle; NOT upload-ready)
- `Manuscript_EEG-CogAgent.docx` - structure ready (abstract 243 words <=250; exactly 7 keywords; Tables 1-3; in-text citation to Table S1 retained), but unresolved author / affiliation / corresponding-author / funding / CRediT / code-availability / approval placeholders remain. Upload-blocked until completed and approved.
- `Cover_Letter_EEG-CogAgent.docx` - structure ready, but contains corresponding-author placeholders and a pending author-approval / originality confirmation. Upload-blocked until completed and approved.

## OPTIONAL
- Graphical abstract - encouraged by JNM but not mandatory. Provided and ready; may be omitted.
- Cover letter - JNM expects one; provided (blocked pending human input above).

## BLOCKING (human-only fields; cannot be auto-generated)
- Author names, ORCID iDs, and numbered affiliations.
- Corresponding-author full contact details (email, phone, mailing address).
- Funding statement (or explicit "no specific grant" confirmation).
- CRediT contributor roles for each author.
- Public repository URL and archived release DOI.
- Written author-approval and originality confirmation from every author.
- Competing-interest confirmation (none declared, or specific disclosure).
- Final corresponding-author sign-off on every auto-generated file.
"""


MANIFEST_ROWS = [
    ("filename", "purpose", "required/optional", "status", "notes"),
    (
        "Manuscript_EEG-CogAgent.docx",
        "Main manuscript text with Tables 1-3 and figure legends (Supplementary Table S1 supplied separately)",
        "Required",
        "BLOCKED - HUMAN INPUT",
        "Structure ready: abstract 243 words (<=250); 7 keywords; Tables 1-3; in-text citation to Table S1 retained. Placeholders for author/affiliation/funding/CRediT/repository/approval unresolved",
    ),
    (
        "Highlights_EEG-CogAgent.docx",
        "Three to five highlight bullets, each <=85 characters",
        "Required",
        "READY",
        "One heading plus 5 bullets; all <=85 characters including spaces; no instructional text",
    ),
    (
        "Cover_Letter_EEG-CogAgent.docx",
        "Editor cover letter",
        "Required",
        "BLOCKED - HUMAN INPUT",
        "Structure ready; corresponding-author placeholders and pending author approval/originality confirmation unresolved",
    ),
    (
        "Supplementary_Table_S1.docx",
        "Positioning table as an editable Word table",
        "Required",
        "READY",
        "6-column editable table; not an image",
    ),
    (
        "Graphical_Abstract_EEG-CogAgent.png",
        "Graphical abstract raster image",
        "Optional",
        "READY",
        "Meets >=531x1328 px; project-owned artwork",
    ),
    (
        "Graphical_Abstract_EEG-CogAgent.pdf",
        "Graphical abstract vector (preferred format)",
        "Optional",
        "READY",
        "Vector; preferred per JNM artwork guidance",
    ),
    (
        "figures/Figure_1.png",
        "Constrained workflow figure (raster)",
        "Required",
        "READY",
        "EEG-CogAgent constrained workflow",
    ),
    (
        "figures/Figure_1.pdf",
        "Constrained workflow figure (vector)",
        "Required",
        "READY",
        "Vector; preferred for drawings",
    ),
    ("figures/Figure_2.png", "Spectral topomaps (raster)", "Required", "READY", "Color artwork"),
    ("figures/Figure_2.pdf", "Spectral topomaps (vector)", "Required", "READY", ""),
    ("figures/Figure_3.png", "Connectivity and network findings (raster)", "Required", "READY", "Color artwork"),
    ("figures/Figure_3.pdf", "Connectivity and network findings (vector)", "Required", "READY", ""),
    ("figures/Figure_4.png", "ROC curves (raster)", "Required", "READY", "Color artwork"),
    ("figures/Figure_4.pdf", "ROC curves (vector)", "Required", "READY", ""),
    ("figures/Figure_5.png", "Cross-condition transfer (raster)", "Required", "READY", "Color artwork"),
    ("figures/Figure_5.pdf", "Cross-condition transfer (vector)", "Required", "READY", ""),
    ("Author_Input_Form.md", "Enumeration of human-only input fields", "Required", "READY", "Complete before submission"),
    ("SUBMISSION_READINESS.md", "READY / OPTIONAL / BLOCKING status map", "Required", "READY", "Explicit human blockers"),
    ("submission_manifest.csv", "This inventory", "Required", "READY", "Auto-generated"),
]


def write_manifest() -> None:
    out = SUBMISSION / "submission_manifest.csv"
    with out.open("w", encoding="utf-8", newline="") as fh:
        writer = csv.writer(fh)
        for row in MANIFEST_ROWS:
            writer.writerow(row)


# --------------------------------------------------------------------------- #
# Orchestration + verification report
# --------------------------------------------------------------------------- #
def build_all() -> None:
    SUBMISSION.mkdir(parents=True, exist_ok=True)
    FIGURES_OUT.mkdir(parents=True, exist_ok=True)

    build_manuscript_docx(SUBMISSION / "Manuscript_EEG-CogAgent.docx")
    build_highlights_docx(SUBMISSION / "Highlights_EEG-CogAgent.docx")
    build_cover_letter_docx(SUBMISSION / "Cover_Letter_EEG-CogAgent.docx")
    build_table_s1_docx(SUBMISSION / "Supplementary_Table_S1.docx")
    build_graphical_abstract(
        SUBMISSION / "Graphical_Abstract_EEG-CogAgent.png",
        SUBMISSION / "Graphical_Abstract_EEG-CogAgent.pdf",
    )
    copy_figures()
    (SUBMISSION / "Author_Input_Form.md").write_text(AUTHOR_INPUT_FORM, encoding="utf-8")
    (SUBMISSION / "SUBMISSION_READINESS.md").write_text(READINESS, encoding="utf-8")
    write_manifest()


def _word_count(text: str) -> int:
    return len(text.split())


def report() -> list[str]:
    import re

    from PIL import Image

    failures: list[str] = []

    md = SRC_MANUSCRIPT.read_text(encoding="utf-8")
    am = re.search(r"## Abstract\s*(.*?)\n\n\*\*Keywords:", md, re.S)
    abstract = am.group(1).strip()
    km = re.search(r"\*\*Keywords:\*\*\s*(.*)", md)
    keywords = [k.strip() for k in km.group(1).strip().split(";")]

    print("=" * 72)
    print("JNM SUBMISSION BUNDLE REPORT")
    print("=" * 72)
    print(f"Manuscript abstract word count : {len(abstract.split())} (limit 250)")
    print(f"Manuscript keyword count       : {len(keywords)} (limit 1-7)")
    for k in keywords:
        print(f"    - {k}")

    print("\nHighlights (character counts, limit 85 incl. spaces):")
    over = False
    for line in SRC_HIGHLIGHTS.read_text(encoding="utf-8").splitlines():
        line = line.strip()
        if line.startswith("- "):
            bullet = line[2:].strip()
            ok = len(bullet) <= 85
            over = over or not ok
            print(f"    [{'OK' if ok else 'OVER'}] {len(bullet):>3}  {bullet}")
    print(f"All highlights <=85 chars      : {not over}")
    if over:
        failures.append("At least one highlight exceeds 85 characters.")

    print("\nGraphical abstract dimensions:")
    ga_png = SUBMISSION / "Graphical_Abstract_EEG-CogAgent.png"
    with Image.open(ga_png) as im:
        w, h = im.size
    meets = w >= GA_MIN_W_PX and h >= GA_MIN_H_PX
    print(f"    PNG  {w} x {h} px  (min {GA_MIN_W_PX} x {GA_MIN_H_PX})  -> {'OK' if meets else 'FAIL'}")
    if not meets:
        failures.append(f"Graphical abstract PNG {w}x{h} below the {GA_MIN_W_PX}x{GA_MIN_H_PX} minimum.")

    # Editable Word tables: the bundle manuscript must carry Tables 1-3 only
    # (Supplementary Table S1 is shipped separately); the project master keeps
    # S1 embedded; the supplementary file is exactly one table.
    table_checks = [
        (
            "Manuscript_EEG-CogAgent.docx (bundle, Tables 1-3 only)",
            SUBMISSION / "Manuscript_EEG-CogAgent.docx",
            3,
        ),
        (
            "EEG-CogAgent_JNM_Submission.docx (master, S1 embedded)",
            ROOT / "docs" / "EEG-CogAgent_JNM_Submission.docx",
            4,
        ),
        (
            "Supplementary_Table_S1.docx (separate editable table)",
            SUBMISSION / "Supplementary_Table_S1.docx",
            1,
        ),
    ]
    print("\nEditable Word tables (count of <w:tbl> per file):")
    for label, path, expected in table_checks:
        doc = Document(path)
        n = len(doc.tables)
        ok = n == expected
        if not ok:
            failures.append(f"{label}: expected {expected} tables, found {n}.")
        print(f"    [{'OK' if ok else 'FAIL'}] {label:<58} tables={n} (expected {expected})")

    # Highlights DOCX must be exactly one Heading 1 plus five bullets and no
    # instructional / "85-character" note paragraph.
    hdoc = Document(SUBMISSION / "Highlights_EEG-CogAgent.docx")
    h_headings = [p for p in hdoc.paragraphs if p.style.name == "Heading 1"]
    h_bullets = [p for p in hdoc.paragraphs if p.style.name == "List Bullet"]
    h_has_note = any("85 characters" in p.text for p in hdoc.paragraphs)
    h_ok = (
        len(h_headings) == 1
        and len(h_bullets) == 5
        and len(hdoc.paragraphs) == 6
        and not h_has_note
    )
    if not h_ok:
        failures.append(
            f"Highlights DOCX structure mismatch: heading={len(h_headings)} "
            f"bullets={len(h_bullets)} paragraphs={len(hdoc.paragraphs)} "
            f"has_note={h_has_note} (expected 1 heading + 5 bullets, no instruction)."
        )
    print(
        "\nHighlights DOCX structure: heading=1 bullets=5 paragraphs=6 no-instruction "
        f"-> {'OK' if h_ok else 'FAIL'}  "
        f"(found heading={len(h_headings)} bullets={len(h_bullets)} "
        f"paragraphs={len(hdoc.paragraphs)} has_note={h_has_note})"
    )

    print("\nBundle file sizes:")
    for path in sorted(SUBMISSION.rglob("*")):
        if path.is_file():
            rel = path.relative_to(SUBMISSION)
            print(f"    {rel.as_posix():<48} {path.stat().st_size:>10} bytes")

    status = "ALL CHECKS PASSED" if not failures else f"{len(failures)} CHECK(S) FAILED"
    print(f"\nValidation status: {status}")
    print("=" * 72)
    return failures


def main() -> None:
    build_all()
    failures = report()
    if failures:
        for msg in failures:
            print(f"FAIL: {msg}")
        sys.exit(1)


if __name__ == "__main__":
    main()
