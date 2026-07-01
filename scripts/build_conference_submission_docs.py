from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_ABSTRACT = ROOT / "docs" / "Conference_Abstract_English.docx"
OUT_ATTACHMENT = ROOT / "docs" / "Figure_Attachment_English.docx"

FIG_DIR = ROOT / "results" / "ds004504_minimal" / "figures"
FIG_DIR_XFER = ROOT / "results" / "ds006036_cross_condition" / "figures"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
GRAY = RGBColor(90, 90, 90)
LIGHT_FILL = "F4F6F9"


TITLE = (
    "EEG-CogAgent: An Auditable Language-Model Agent for Reproducible "
    "Dementia EEG Biomarker Analysis"
)


ABSTRACT_SECTIONS = [
    (
        "Objective",
        "Public electroencephalography (EEG) datasets provide an important "
        "basis for reproducible biomarker research in cognitive disorders, "
        "but converting Brain Imaging Data Structure recordings into "
        "traceable analyses and manuscript-ready reports remains labor "
        "intensive. This study aimed to develop and validate EEG-CogAgent, "
        "a lightweight large language model (LLM)-assisted skill for "
        "automated resting-state EEG analysis in dementia. The goal was not "
        "to let an LLM make clinical diagnoses, but to use it as an auditable "
        "analysis coordinator that plans, checks, and summarizes deterministic "
        "EEG workflows."
    ),
    (
        "Methods",
        "We implemented EEG-CogAgent as a configuration-driven agent that "
        "orchestrates Python modules for BIDS loading, preprocessing, feature "
        "extraction, statistical testing, machine-learning validation, "
        "visualization, artifact auditing, and report drafting. The LLM was "
        "restricted to workflow planning, tool selection, output checking, "
        "and narrative synthesis; it did not inspect raw EEG for diagnosis "
        "or alter numerical results. The framework was evaluated using "
        "OpenNeuro ds004504, an eyes-closed resting-state EEG dataset "
        "including 36 patients with Alzheimer's disease (AD), 23 patients "
        "with frontotemporal dementia (FTD), and 29 healthy controls (HC), "
        "recorded from 19 scalp channels at 500 Hz. Signals were band-pass "
        "filtered from 1 to 45 Hz, notch-filtered at 50 Hz, re-referenced to "
        "the average reference, segmented into 4-s epochs, and screened using "
        "a 150 microvolt peak-to-peak rejection rule. The feature set included "
        "frequency-band power, theta/alpha and delta/alpha ratios, regional "
        "power summaries, coherence, phase-lag index, weighted phase-lag "
        "index, clustering coefficient, global efficiency, and mean node "
        "strength. Statistical analyses used group tests with false discovery "
        "rate correction and covariate-adjusted models. Logistic regression, "
        "support vector machine, and random forest classifiers were evaluated "
        "with participant-level nested cross-validation. A paired "
        "condition-transfer analysis used ds006036 photomark recordings from "
        "the same cohort as an acquisition-condition robustness test rather "
        "than as independent external validation."
    ),
    (
        "Results",
        "All 88 participants in the primary dataset were processed, and the "
        "workflow passed 22 deterministic audit checks covering participant "
        "uniqueness, cohort counts, numerical ranges, out-of-fold prediction "
        "coverage, probability validity, and required output artifacts. After "
        "adjustment for age, gender, and retained epoch count, 104 of 137 "
        "spectral features differed among diagnostic groups after false "
        "discovery rate correction. The strongest adjusted marker was the "
        "temporal theta/alpha ratio (Wald chi-square = 55.16, q = "
        "1.44 x 10^-10), followed by global and regional theta/alpha ratios. "
        "Pairwise analyses showed extensive AD-versus-HC and FTD-versus-HC "
        "effects, but no spectral or connectivity feature remained "
        "significant for AD versus FTD after correction. Connectivity results "
        "were directionally consistent with disease-associated slowing and "
        "altered network organization, including lower alpha coherence and "
        "higher theta weighted phase-lag strength in dementia groups relative "
        "to controls. In three-class classification, the best model was the "
        "support vector machine, with accuracy 0.614, balanced accuracy "
        "0.593, and multiclass AUC 0.727. Disease-versus-control tasks "
        "performed better: AD versus HC reached balanced accuracy 0.837, and "
        "FTD versus HC reached balanced accuracy 0.744. AD-versus-FTD "
        "classification was weaker and did not support a reliable differential "
        "diagnosis claim. In participant-disjoint paired condition transfer, "
        "the SVM achieved photomark balanced accuracy 0.712 and AUC 0.794, "
        "indicating within-cohort condition transfer rather than external "
        "generalization."
    ),
    (
        "Conclusion",
        "EEG-CogAgent demonstrates that an LLM-assisted skill can coordinate "
        "a reproducible, auditable EEG biomarker workflow while preserving a "
        "clear boundary between language-model orchestration and deterministic "
        "scientific computation. In a public dementia EEG cohort, the agent "
        "recovered interpretable slowing-related biomarkers and generated "
        "validated analysis artifacts, but weak AD-versus-FTD separation and "
        "shared-cohort condition transfer indicate that the present evidence "
        "supports a research-assistance framework, not an autonomous clinical "
        "diagnostic system. Future work should lock the configuration before "
        "multisite external validation and compare agent-assisted analysis "
        "with conventional manual workflows."
    ),
]


FIGURES = [
    {
        "title": "Figure 1. EEG-CogAgent workflow.",
        "paths": [FIG_DIR / "figure1_workflow.png"],
        "width": 6.2,
        "caption": (
            "The skill reads BIDS EEG data, calls deterministic MNE-Python-based "
            "modules, writes reusable analysis artifacts, audits the outputs, "
            "and then generates a constrained report draft."
        ),
        "alt": "Workflow diagram showing BIDS input, preprocessing, feature extraction, statistics, machine learning, audit, and report generation.",
    },
    {
        "title": "Figure 2. Spectral topographies.",
        "paths": [
            FIG_DIR / "topomap_delta.png",
            FIG_DIR / "topomap_theta.png",
            FIG_DIR / "topomap_alpha.png",
            FIG_DIR / "topomap_beta.png",
            FIG_DIR / "topomap_gamma.png",
        ],
        "width": 3.05,
        "caption": (
            "Band-power scalp maps summarize interpretable resting-state EEG "
            "biomarkers across delta, theta, alpha, beta, and gamma bands."
        ),
        "alt": "Five scalp topography panels for delta, theta, alpha, beta, and gamma band power.",
    },
    {
        "title": "Figure 3. Connectivity and graph features.",
        "paths": [FIG_DIR / "figure3_connectivity.png"],
        "width": 6.2,
        "caption": (
            "Connectivity summaries emphasize altered theta-band delayed phase "
            "coupling and reduced alpha coherence in dementia groups relative "
            "to controls."
        ),
        "alt": "Connectivity figure comparing theta weighted phase-lag networks and group-level graph metrics.",
    },
    {
        "title": "Figure 4. Nested cross-validation receiver operating curves.",
        "paths": [FIG_DIR / "figure4_roc.png"],
        "width": 6.2,
        "caption": (
            "Three-class classification was moderate, with stronger "
            "disease-versus-control separation than AD-versus-FTD separation."
        ),
        "alt": "Receiver operating characteristic curves for logistic regression, support vector machine, and random forest models.",
    },
    {
        "title": "Figure 5. Paired condition-transfer validation.",
        "paths": [FIG_DIR_XFER / "figure5_cross_condition.png"],
        "width": 6.2,
        "caption": (
            "Eyes-closed-trained models transferred to held-out photomark "
            "recordings from the same cohort; this supports condition robustness "
            "but not external validation."
        ),
        "alt": "Bar and point plot showing paired eyes-closed to photomark condition-transfer balanced accuracy and AUC.",
    },
]


def set_run_font(run, name="Calibri", size=11, bold=None, italic=None, color=None):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rpr.get_or_add_rFonts().set(qn("w:ascii"), name)
    rpr.get_or_add_rFonts().set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_doc_defaults(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run_font(run, size=9, color=GRAY)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def set_footer(doc):
    footer = doc.sections[0].footer
    if footer.paragraphs:
        footer.paragraphs[0].text = ""
        add_page_number(footer.paragraphs[0])


def add_title_block(doc, subtitle):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(TITLE)
    set_run_font(r, size=20, bold=True, color=DARK_BLUE)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run(subtitle)
    set_run_font(r, size=11, italic=True, color=GRAY)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    for label, value in [
        ("Authors", "[Author names]"),
        ("Affiliations", "[Affiliations]"),
        ("Corresponding author", "[Name and email]"),
    ]:
        rr = p.add_run(f"{label}: ")
        set_run_font(rr, bold=True)
        rr = p.add_run(f"{value}\n")
        set_run_font(rr)


def add_callout(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0)
    p.paragraph_format.right_indent = Inches(0)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(14)
    p.paragraph_format.line_spacing = 1.10
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), LIGHT_FILL)
    p_pr.append(shd)
    r = p.add_run(text)
    set_run_font(r, size=10.5, color=DARK_BLUE)


def set_last_picture_alt(descr):
    # python-docx does not expose alt text directly; patch the most recent docPr.
    # The package-level query is safe here because each insertion happens serially.
    return descr


def patch_inline_alt(doc, descrs):
    doc_prs = doc._element.xpath(".//wp:docPr")
    for node, descr in zip(doc_prs, descrs):
        node.set("descr", descr)


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_run_font(r, size=9.5, italic=True, color=GRAY)


def build_abstract():
    doc = Document()
    set_doc_defaults(doc)
    set_footer(doc)
    add_title_block(doc, "English conference abstract for online submission")

    add_callout(
        doc,
        "Submission note: the abstract body is structured as Objective, Methods, "
        "Results, and Conclusion, and intentionally contains no figures or tables.",
    )

    doc.add_heading("Abstract", level=1)
    for heading, text in ABSTRACT_SECTIONS:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(f"{heading}: ")
        set_run_font(r, bold=True)
        r = p.add_run(text)
        set_run_font(r)

    p = doc.add_paragraph()
    r = p.add_run("Keywords: ")
    set_run_font(r, bold=True)
    r = p.add_run(
        "electroencephalography; Alzheimer's disease; frontotemporal dementia; "
        "large language model; agent; biomarker; reproducibility"
    )
    set_run_font(r)

    doc.save(OUT_ABSTRACT)


def build_attachment():
    doc = Document()
    set_doc_defaults(doc)
    set_footer(doc)
    add_title_block(doc, "Figure attachment for conference submission")

    add_callout(
        doc,
        "This attachment provides visual material only. The online abstract should "
        "remain text-only according to the conference requirement."
    )

    doc.add_heading("Key message", level=1)
    bullets = [
        "EEG-CogAgent coordinates deterministic EEG analysis rather than making autonomous diagnoses.",
        "The dementia case study recovered interpretable slowing-related biomarkers in OpenNeuro ds004504.",
        "Classification was strongest for disease-versus-control separation and weaker for AD-versus-FTD differential diagnosis.",
        "The paired ds006036 experiment indicates within-cohort condition transfer, not independent external validation.",
    ]
    for item in bullets:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(item)
        set_run_font(r)

    alt_texts = []
    for idx, fig in enumerate(FIGURES, start=1):
        if idx > 1:
            doc.add_section(WD_SECTION.NEW_PAGE)
        doc.add_heading(fig["title"], level=1)
        paths = fig["paths"]
        if len(paths) == 1:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(str(paths[0]), width=Inches(fig["width"]))
            alt_texts.append(fig["alt"])
        else:
            for row_start in (0, 2, 4):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                row_paths = paths[row_start : row_start + 2]
                for path_i, path in enumerate(row_paths):
                    p.add_run().add_picture(str(path), width=Inches(fig["width"]))
                    alt_texts.append(fig["alt"])
                    if path_i < len(row_paths) - 1:
                        p.add_run("  ")
        add_caption(doc, fig["caption"])

    patch_inline_alt(doc, alt_texts)
    doc.save(OUT_ATTACHMENT)


def main():
    build_abstract()
    build_attachment()
    print(OUT_ABSTRACT)
    print(OUT_ATTACHMENT)


if __name__ == "__main__":
    main()
