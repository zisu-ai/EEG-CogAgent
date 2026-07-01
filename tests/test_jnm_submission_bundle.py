"""Integration tests for the JNM submission bundle builder.

These rebuild the project master manuscript and the submission bundle, then
verify the structural invariants Codex's review required. The builders target
``docs/`` and ``submission/JNM/`` by design and are idempotent, so the tests
write to those real locations.
"""

from __future__ import annotations

import csv
import sys
from pathlib import Path

import pytest
from docx import Document

ROOT = Path(__file__).resolve().parents[1]
SCRIPTS = ROOT / "scripts"
if str(SCRIPTS) not in sys.path:
    sys.path.insert(0, str(SCRIPTS))

import build_jnm_manuscript_docx as jnm  # noqa: E402
import build_jnm_submission_bundle as bundle  # noqa: E402

SUBMISSION = ROOT / "submission" / "JNM"
MASTER = ROOT / "docs" / "EEG-CogAgent_JNM_Submission.docx"


@pytest.fixture(scope="module", autouse=True)
def _rebuild():
    """Rebuild master (with S1) and the bundle (without S1) once for the module."""
    jnm.build()        # default: master keeps Supplementary Table S1 embedded
    bundle.build_all()  # bundle manuscript omits S1; ships it as a separate file
    yield


def _tables(path: Path) -> int:
    return len(Document(path).tables)


def test_bundle_manuscript_has_only_tables_1_to_3():
    assert _tables(SUBMISSION / "Manuscript_EEG-CogAgent.docx") == 3


def test_master_manuscript_keeps_embedded_s1():
    assert _tables(MASTER) == 4


def test_supplementary_table_s1_is_one_table():
    assert _tables(SUBMISSION / "Supplementary_Table_S1.docx") == 1


def test_bundle_manuscript_retains_intext_s1_citation():
    text = "\n".join(
        p.text for p in Document(SUBMISSION / "Manuscript_EEG-CogAgent.docx").paragraphs
    )
    assert "Table S1" in text


def test_highlights_docx_is_one_heading_plus_five_bullets():
    doc = Document(SUBMISSION / "Highlights_EEG-CogAgent.docx")
    headings = [p for p in doc.paragraphs if p.style.name == "Heading 1"]
    bullets = [p for p in doc.paragraphs if p.style.name == "List Bullet"]
    assert len(doc.paragraphs) == 6
    assert len(headings) == 1
    assert len(bullets) == 5
    assert not any("85 characters" in p.text for p in doc.paragraphs)


def test_cover_letter_heading_is_cover_letter():
    doc = Document(SUBMISSION / "Cover_Letter_EEG-CogAgent.docx")
    headings = [p.text for p in doc.paragraphs if p.style.name == "Heading 1"]
    assert headings == ["Cover Letter"]


def test_cover_letter_keeps_pending_marker_and_placeholders():
    text = "\n".join(p.text for p in Document(SUBMISSION / "Cover_Letter_EEG-CogAgent.docx").paragraphs)
    assert "[Pending:" in text
    assert "[Corresponding author name]" in text


def test_manifest_marks_manuscript_and_cover_letter_blocked():
    with (SUBMISSION / "submission_manifest.csv").open(encoding="utf-8") as fh:
        rows = list(csv.reader(fh))
    status_idx = rows[0].index("status")
    by_file = {r[0]: r[status_idx] for r in rows[1:]}
    assert by_file["Manuscript_EEG-CogAgent.docx"] == "BLOCKED - HUMAN INPUT"
    assert by_file["Cover_Letter_EEG-CogAgent.docx"] == "BLOCKED - HUMAN INPUT"
    assert by_file["Highlights_EEG-CogAgent.docx"] == "READY"
    assert by_file["Supplementary_Table_S1.docx"] == "READY"


def test_readiness_states_blocked_human_input():
    text = (SUBMISSION / "SUBMISSION_READINESS.md").read_text(encoding="utf-8")
    assert "BLOCKED - HUMAN INPUT" in text
    assert "NOT upload-ready" in text
    assert "Manuscript_EEG-CogAgent.docx" in text
    assert "Cover_Letter_EEG-CogAgent.docx" in text


def test_report_returns_no_failures():
    assert bundle.report() == []
